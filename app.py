import csv
import os
import pickle
from urllib.parse import unquote

from PyPDF2 import PdfReader
from docx import Document
import docx2txt
import requests
from bs4 import BeautifulSoup
from PIL import Image
import pytesseract
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from langchain_unstructured import UnstructuredLoader
from unstructured.cleaners.core import clean_extra_whitespace
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from moviepy import VideoFileClip
import whisper
from langchain_core.prompts import ChatPromptTemplate
from langchain_mistralai import ChatMistralAI
import google.generativeai as genai
import re
from flask import Flask, request, jsonify, render_template, send_from_directory
from werkzeug.utils import secure_filename

import base64
from io import BytesIO

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
VECTOR_FOLDER = 'vector_stores'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'csv', 'txt'}

os.environ['FAISS_OPT_LEVEL'] = 'avx512'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50 MB
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['VECTOR_FOLDER'] = VECTOR_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(VECTOR_FOLDER, exist_ok=True)


class WordImageTextExtractor:
    """
    Extracts text and images from a Word file.
    Uses OCR to extract text from images.
    Returns a single combined text output.
    """

    def __init__(self, docx_path, images_folder):
        self.docx_path = docx_path
        self.images_folder = images_folder
        os.makedirs(self.images_folder, exist_ok=True)

    def extract_text_and_images(self):
        """
        Extracts text and images from the Word document.
        Runs OCR on images and combines all extracted text.
        :return: Combined text with labels.
        """
        # Extract text and images
        text = docx2txt.process(self.docx_path, self.images_folder)
        combined_text = f"Extracted Document Text:\n{text}\n\n"

        # Run OCR on extracted images
        image_files = sorted(os.listdir(self.images_folder))  # Sort to maintain order
        for img_file in image_files:
            img_path = os.path.join(self.images_folder, img_file)

            try:
                img_text = pytesseract.image_to_string(Image.open(img_path))
                combined_text += f"Images Text ({img_file}):\n{img_text}\n\n"
            except Exception as e:
                combined_text += f"Error processing {img_file}: {e}\n\n"

        return combined_text.strip()


class DocumentLoader:
    def __init__(self, file_path):
        self.file_path = file_path
        self.filename = os.path.basename(file_path)

    def read_pdf_text(self):
        """Extracts text with page numbers from PDF"""
        documents = []
        try:
            with open(self.file_path, 'rb') as pdf_file:
                pdf_reader = PdfReader(pdf_file, strict=False)
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    text = page.extract_text() or ""
                    if not text.strip():
                        raise ValueError("Empty text extracted from PDF pages.")
                    doc = LangchainDocument(
                        page_content=text.strip(),
                        metadata={'page_no': page_num, 'filename': self.filename}
                    )
                    documents.append(doc)
            return documents
        except Exception:
            return self.fallback_unstructured_extraction()

    def read_docx_text(self):
        """Extracts text from DOCX, falls back to OCR if text is empty"""
        documents = []
        try:
            doc = Document(self.file_path)
            for para_num, para in enumerate(doc.paragraphs, start=1):
                if para.text.strip():
                    doc = LangchainDocument(
                        page_content=para.text,
                        metadata={'section_no': para_num, 'filename': self.filename}
                    )
                    documents.append(doc)

            if not documents:  # If no text was extracted, use OCR
                images_folder = f"{self.file_path}_images"
                text_from_images = WordImageTextExtractor(self.file_path, images_folder).extract_text_and_images()
                return [LangchainDocument(
                    page_content=text_from_images,
                    metadata={'filename': self.filename}
                )]
            return documents
        except Exception:
            return self.fallback_unstructured_extraction()

    def read_csv_text(self):
        """Extracts text with row numbers from CSV"""
        documents = []
        try:
            with open(self.file_path, newline='', encoding='utf-8') as csvfile:
                reader = csv.reader(csvfile)
                for row_num, row in enumerate(reader, start=1):
                    doc = LangchainDocument(
                        page_content=" ".join(row),
                        metadata={'row_no': row_num, 'filename': self.filename}
                    )
                    documents.append(doc)
            return documents
        except Exception:
            return self.fallback_unstructured_extraction()

    def read_txt_text(self):
        """Extracts text from a TXT file"""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as txt_file:
                full_text = txt_file.read()
            return [LangchainDocument(
                page_content=full_text,
                metadata={'page_no': 1, 'filename': self.filename}
            )]
        except Exception:
            return self.fallback_unstructured_extraction()

    def fallback_unstructured_extraction(self):
        """Fallback to UnstructuredLoader if other methods fail"""
        try:
            loader = UnstructuredLoader(self.file_path, post_processors=[clean_extra_whitespace])
            return loader.load()
        except Exception as e:
            raise ValueError(f"Failed to process {self.file_path} with both structured and unstructured methods. Error: {e}")
    
    def read_url_text(self,url):
        """Extracts text from URLs and generates a filename."""
        documents = []
        try:
            response = requests.get(url)
            response.raise_for_status()
            content = response.text
            # Extract text from the webpage
            soup = BeautifulSoup(content, 'html.parser')
            text = soup.get_text(separator='\n').strip()

            self.file_path = soup.title.get_text() + '.html'
            # Create a LangchainDocument
            doc = LangchainDocument(
                page_content=text,
                metadata={
                    'url': url,
                    'filename': self.file_path
                }
            )
            documents.append(doc)
        except Exception as e:
            print(f"Error reading URL: {e}")
            raise ValueError("Failed to process URL. It may be invalid or inaccessible.")
        return documents
    
    def load_document(self):
        """Loads documents with format-specific metadata"""
        if self.file_path.endswith('.pdf'):
            return self.read_pdf_text()
        elif self.file_path.endswith('.docx'):
            return self.read_docx_text()
        elif self.file_path.endswith('.csv'):
            return self.read_csv_text()
        elif self.file_path.endswith('.txt'):
            return self.read_txt_text()
        else:
            raise ValueError("Unsupported file format")

    def chunk_text(self, documents, chunk_size=10000, chunk_overlap=1000):
        """Splits documents while preserving metadata"""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n"]
        )
        chunked_documents = []
        for doc in documents:
            split_texts = text_splitter.split_text(doc.page_content)
            for chunk in split_texts:
                chunked_documents.append(
                    LangchainDocument(
                        page_content=chunk,
                        metadata=doc.metadata
                    )
                )
        return chunked_documents

class VectorEmbeddingRetrieval:
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):

        self.embeddings = HuggingFaceEmbeddings(model_name=model_name)

    def create_vector_store(self, chunks):
        """Create FAISS index with enhanced metadata handling"""
        return FAISS.from_documents(
            documents=chunks,
            embedding=self.embeddings
        )

class VideoProcessor:
    """
    A class to handle video-related operations, such as extracting audio from video files.
    """

    def __init__(self):
        pass

    def extract_audio(self, video_path,audio_path="extracted_audio.mp3"):
        """Extracts audio from the video file and saves it as an audio file."""
        try:
            # Load the video file
            with VideoFileClip(video_path) as video_clip:
                
                # Extract audio from video
                audio_clip = video_clip.audio
                
                if audio_clip:
                    # Save the audio file
                    audio_clip.write_audiofile(audio_path)
                    print(f"Successfully converted {video_path} to {audio_path}")
                    return audio_path
                else:
                    print("No audio stream found in the video file.")
                    return None
                    
        except Exception as e:
            print(f"Error occurred: {str(e)}")
            return None


class AudioProcessor:
    """A class to process audio files and transcribe speech using Whisper."""

    def __init__(self, model_size="tiny.en"):
        """Initialize the Whisper model with the specified size."""
        self.model = whisper.load_model(model_size)

    def transcribe_audio(self, audio_path):
        """Transcribes the given audio file and returns the text."""
        try:
            result = self.model.transcribe(audio_path)
            return result['text']
        except Exception as e:
            print(f"Error occurred during transcription: {str(e)}")
            return None

# HF_TOKEN = os.getenv("HF_TOKEN")
MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")
os.environ["MISTRAL_API_KEY"] = MISTRAL_API_KEY
chat_history = []
# from langchain import HuggingFaceEndpoint
class QuestionAnswering:
  def __init__(self):

    self.llm = ChatMistralAI(
    model="mistral-large-latest",
    temperature=0,
    max_retries=2,
    # other params...
    )
    # self.llm = HuggingFaceEndpoint(
    #     repo_id="mistralai/Mistral-7B-Instruct-v0.2",
    #     huggingfacehub_api_token=HF_TOKEN,
    #     temperature=0.5,
    #     max_new_tokens=512,
    #     task="text-generation"  # Explicitly specify the task type
    # )

  def answer_question(self,contexts, question, chat_history=None):

    # If no chat history is provided, initialize an empty list
    if chat_history is None:
        chat_history = []

    # Create a custom prompt template with system and user roles
    system_prompt = """You are a helpful AI assistant.You are a Document Question Answer Assistant.  
    Use the following pieces of retrieved context to answer the question.
    Answer the question based on the context provided. 
    Reply the answer like conversation manner.
    Keep the answer concise and relevant.
    Only Return The Answer.

    Chat History:
    {chat_history}

    Context:
    {context}"""

    user_prompt = """Question: {input}
    Helpful Answer:"""

    # Create a chat prompt template
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        ("user", user_prompt)
    ])

    rag_chain = prompt | self.llm

    answer = rag_chain.invoke({
        "input": question,
        "chat_history": chat_history,
        "context": contexts
    })

    #return answer # for langchain-mistral
    return answer.content
    
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
class MultiModalSupport:
    def __init__(self):
        # Configure Gemini API
        genai.configure(api_key=GEMINI_API_KEY) # Replace with your actual API key
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    def encode_image_to_base64(self,image_path):
        """Encodes an image to base64."""
        try:
            img_pil = Image.open(image_path)
            img_pil = img_pil.convert('RGB')
            buffered = BytesIO()
            img_pil.save(buffered, format="JPEG")  # You can use other formats like PNG
            img_str = base64.b64encode(buffered.getvalue()).decode()
            return img_str
        except Exception as e:
            print(f"Error encoding image: {e}")
            return None

    def get_image_description(self,image_path):
        """Gets the image description from the Gemini model."""
        try:
            base64_image = self.encode_image_to_base64(image_path)
            if base64_image:
                image_parts = [
                    {
                        "mime_type": "image/jpeg",  # Or "image/png"
                        "data": base64_image
                    },
                    {"text": "Analyze the image to Give me the Detailed Information About the Image:"}, # Prompt to get the information.
                ]
                response = self.model.generate_content(image_parts)
                print(response)
                print("Image to text conversion successful!")
                return response.text
            else:
                return "Image encoding failed."
        except Exception as e:
            print(f"Error getting image description: {e}")
            return "An error occurred."

vector_retrieval = VectorEmbeddingRetrieval()
qa_system = QuestionAnswering()

video_processor = VideoProcessor()
audio_processor = AudioProcessor()
image_processor = MultiModalSupport()


def write_txt_file( file_path,transcribed_text):
    try:
        with open(file_path + ".txt", "w", encoding="utf-8") as file:
            file.write(transcribed_text)

        file_path = file_path + ".txt"
        print(f"Transcribed text saved to {file_path}")

        return file_path
    except IOError as e:
        print(f"Error saving transcribed text: {e}")

def video_processing(file_path):
    try:
        audio_path = video_processor.extract_audio(file_path)

        if audio_path:
            try:
                transcribed_text = audio_processor.transcribe_audio(audio_path)

                if transcribed_text:
                    # Save the transcribed text to a file
                    file_path = write_txt_file(file_path,transcribed_text)
            except Exception as e:
                print(f"Error processing audio: {e}")
    except Exception as e:
        print(f"Unexpected error: {e}")

    return file_path

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].strip().lower() in ALLOWED_EXTENSIONS

def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

def get_vector_path(filename):
    return os.path.join(app.config['VECTOR_FOLDER'], f"{filename}.pkl")

@app.route('/')
def home():
    return render_template('upload.html')

@app.route('/chat')
def chat():
    return render_template('chat.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    print(f"Received file: {file.filename}")


    try:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(file_path):
            print(f"File '{filename}' already exists. Skipping upload.")
        else:
            file.save(file_path)
            print(f"File '{filename}' saved successfully.")

        if filename.lower().endswith(".mp4"):
            file_path = video_processing(file_path)

        elif filename.lower().endswith((".mp3", ".wav")):
            transcribed_text = audio_processor.transcribe_audio(file_path)
            file_path = write_txt_file(file_path, transcribed_text)

        elif filename.lower().endswith((".jpg", ".jpeg", ".png")):
            img_text = image_processor.get_image_description(file_path)
            file_path = write_txt_file(file_path, img_text)


        document_loader = DocumentLoader(file_path)
        documents = document_loader.load_document()
        chunks = document_loader.chunk_text(documents)

        vector_store = vector_retrieval.create_vector_store(chunks)

        vector_path = get_vector_path(filename)
        with open(vector_path, 'wb') as f:
            pickle.dump(vector_store, f)

        return jsonify({"success": True, "message": f"File '{filename}' uploaded and processed successfully"})
    except Exception as e:
        print(f"Error uploading file: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/process_urls', methods=['POST'])
def process_urls():
    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 415

    data = request.get_json()
    print(data)
    urls = data.get('urls')

    all_chunks = []
    webpage_content = ""
    try:
        for idx, url in enumerate(urls):
            # Attempt to fetch and process each URL
            try:
                document_loader = DocumentLoader(f"url{idx + 1}.html")
                documents = document_loader.read_url_text(url)
                # Save HTML content in uploads folder

                html_file_name = sanitize_filename(document_loader.file_path) + ".txt"
                txt_file_path = os.path.join(UPLOAD_FOLDER, html_file_name)
                
                for document in documents:
                    webpage_content += document.page_content + "\n"

                with open(txt_file_path, "w", encoding="utf-8") as html_file:
                    html_file.write(webpage_content)

                chunks = document_loader.chunk_text(documents)

                for chunk in chunks:
                    chunk.metadata['page_no'] = idx + 1

                all_chunks.extend(chunks)
            except Exception as e:
                return jsonify({"error": f"Failed to process URL {url}: {str(e)}"}), 500

        # Create vector store from all chunks
        vector_store = vector_retrieval.create_vector_store(all_chunks)

        # Store the vector store
        for idx in range(len(urls)):
            vector_path = get_vector_path(sanitize_filename(document_loader.file_path))
            with open(vector_path, 'wb') as f:
                pickle.dump(vector_store, f)

        return jsonify({"success": True, "message": "URLs processed and stored successfully."})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@app.route('/ask', methods=['POST'])
def ask_question():
    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 415

    data = request.get_json()
    question = data.get('question')
    filename = data.get('file')

    if not question:
        return jsonify({"error": "Question is required"}), 400
    if not filename:
        return jsonify({"error": "No file selected"}), 400

    vector_path = get_vector_path(filename)
    if not os.path.exists(vector_path):
        return jsonify({"error": "File not found in vector store"}), 404

    try:
        with open(vector_path, 'rb') as f:
            vector_store = pickle.load(f)

        retriever = vector_store.similarity_search(question,k=3)
        answer = qa_system.answer_question(retriever, question, chat_history)
        # answer = "Hello! How Can I Help You"
        chat_history.append((question, answer))
        return jsonify({"answer": answer})
    except Exception as e:
        print(f"Error answering question: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/get_context', methods=['POST'])
def get_context():
    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 415

    data = request.get_json()
    question = data.get('question')
    filename = data.get('file')

    if not question:
        return jsonify({"error": "Question is required"}), 400
    if not filename:
        return jsonify({"error": "No file selected"}), 400

    vector_path = get_vector_path(filename)
    if not os.path.exists(vector_path):
        return jsonify({"error": "File not found in vector store"}), 404

    try:
        with open(vector_path, 'rb') as f:
            vector_store = pickle.load(f)
        
        contexts = vector_store.similarity_search(question,k=3)
        print(contexts)
        serialized_contexts = [
            {
                "page_content": doc.page_content,
                "metadata": doc.metadata
            }
            for doc in contexts
        ]

        return jsonify({
            "contexts": serialized_contexts,
            "page_no": serialized_contexts[0]['metadata'].get('section_no') if serialized_contexts else None,
            "file_name": serialized_contexts[0]['metadata'].get('filename') if serialized_contexts else None
        })
    except Exception as e:
        print(f"Error retrieving context: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/list_files', methods=['GET'])
def list_files():
    files = [f.replace(".pkl", "") for f in os.listdir(app.config['VECTOR_FOLDER']) if f.endswith(".pkl")]
    return jsonify({"files": files})

@app.route('/preview/<filename>')
def preview_file(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404

        document_loader = DocumentLoader(file_path)
        content = document_loader.load_document()

        return jsonify({"content": content[0].page_content}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
from urllib.parse import unquote
@app.route('/api/file/<path:filename>')
def api_file(filename):
    filename = unquote(filename) # decode url encoded filenames.
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route("/delete-file", methods=["DELETE"])
def delete_file():
    """Deletes a file from the vector store folder."""
    filename = request.args.get("filename")
    print(f"Deleting file: {filename}")
    if "%20" in filename:
        filename = filename.replace("%20", " ")

    if not filename:
        return jsonify({"success": False, "error": "Filename not provided"}), 400

    file_path = os.path.join(app.config["VECTOR_FOLDER"], filename +".pkl")
    
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            return jsonify({"success": True, "message": f"Deleted {filename}"})
        except Exception as e:
            return jsonify({"success": False, "error": str(e)}), 500
    else:
        return jsonify({"success": False, "error": "File not found"}), 404
    
if __name__ == '__main__':
    app.run(debug=True, use_reloader=False)
